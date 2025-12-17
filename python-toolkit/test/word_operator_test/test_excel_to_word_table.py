import os
from docx import Document
from openpyxl import Workbook
import sys
sys.path.append('../src')
from word.excel_to_word_table import ExcelToWordTable
# 使用示例函数
def example_usage():
    """使用示例"""
    
    # 示例1: 基本使用
    print("示例1: 基本使用")
    print("-" * 50)
    
    # 创建实例
    tool = ExcelToWordTable(
        excel_path="data/sample.xlsx",  # 你的Excel文件路径
        word_path="output/report.docx"  # 输出Word文件路径
    )
    
    # 从Excel加载数据
    # 从第1个工作表的A1单元格开始读取
    table_data = tool.load_excel(
        sheet_name="Sheet1",  # 工作表名
        start_row=1,         # 起始行
        start_col=1,         # 起始列
        end_row=10,          # 结束行
        end_col=5           # 结束列
    )
    
    print(f"读取到 {len(table_data)} 行数据")
    if table_data:
        print(f"每行 {len(table_data[0])} 列")
        print("前3行数据示例:")
        for i, row in enumerate(table_data[:3]):
            print(f"  行{i+1}: {row[:5]}...")
    
    # 创建Word文档（可以基于模板）
    tool.create_word_document(template_path="templates/template.docx")
    
    # 在指定占位符位置插入表格
    tool.insert_table_at_placeholder(
        table_data=table_data,
        placeholder="{{ table1 }}",  # Word中的占位符
        auto_adjust_columns=True,
        table_style="Table Grid"
    )
    
    # 保存文档
    tool.save_word_document()
    print("表格已插入到Word文档")
    
    '''
    # 示例2: 批量插入多个表格
    print("\n示例2: 批量插入多个表格")
    print("-" * 50)
    
    tool2 = ExcelToWordTable("data/multi_tables.xlsx")
    
    # 加载多个表格
    summary_data = tool2.load_excel("Summary", 1, 1, 5, 4)
    details_data = tool2.load_excel("Details", 1, 1, 20, 6)
    
    # 创建新文档
    tool2.create_word_document()
    
    # 添加标题
    doc = tool2.word_doc
    doc.add_heading('数据报告', 0)
    doc.add_paragraph('以下是数据摘要: {{ summary }}')
    doc.add_paragraph('详细数据如下: {{ details }}')
    
    # 批量替换占位符
    placeholders = {
        "{{ summary }}": summary_data,
        "{{ details }}": details_data
    }
    tool2.replace_multiple_placeholders(placeholders)
    
    # 保存
    tool2.save_word_document("output/multi_table_report.docx")
    print("多表格文档已生成")
    '''

def test_excel_to_word_table():
    print("运行示例程序...")
    print("=" * 60)
    
    # 创建示例数据目录
    os.makedirs("data", exist_ok=True)
    os.makedirs("templates", exist_ok=True)
    os.makedirs("output", exist_ok=True)
    
    # 创建示例Excel文件
    
    
    wb = Workbook()
    ws = wb.active
    ws.title = "Sheet1"
    
    # 添加示例数据
    headers = ["ID", "姓名", "部门", "销售额", "完成率"]
    ws.append(headers)
    
    sample_data = [
        [1, "张三", "销售部", 150000, "85%"],
        [2, "李四", "技术部", 120000, "92%"],
        [3, "王五", "市场部", 180000, "78%"],
        [4, "赵六", "销售部", 210000, "95%"],
        [5, "孙七", "技术部", 95000, "88%"],
    ]
    
    for row in sample_data:
        ws.append(row)
    
    wb.save("data/sample.xlsx")
    print("已创建示例Excel文件: data/sample.xlsx")
    
    # 创建示例Word模板
    doc = Document()
    doc.add_heading('销售报告', 0)
    doc.add_paragraph('以下是从Excel导入的数据表格:')
    doc.add_paragraph('{{ table1 }}')
    doc.add_paragraph('表格结束')
    doc.save("templates/template.docx")
    print("已创建示例Word模板: templates/template.docx")
    
    # 运行示例
    example_usage()
    
    print("\n" + "=" * 60)
    print("示例完成！")
    print("\n也可以使用命令行模式:")
    print("python excel_to_word.py data/sample.xlsx --template templates/template.docx --output output/result.docx")
