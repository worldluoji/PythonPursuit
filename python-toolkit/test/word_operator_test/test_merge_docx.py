from docx import Document
import os
import sys
sys.path.append('../src')
from word.docx_merger import DocxContentMerger

def create_test_documents():
    """创建测试文档"""
    print("创建测试文档...")
    
    # 创建源文档A
    doc_a = Document()
    doc_a.add_heading('源文档A - 测试内容', 0)
    doc_a.add_paragraph('这是从源文档A读取的第一段内容。')
    doc_a.add_paragraph('这是第二段内容，包含一些格式：')
    
    # 添加有格式的段落
    p = doc_a.add_paragraph()
    p.add_run('加粗文本').bold = True
    p.add_run(' 普通文本 ')
    p.add_run('斜体文本').italic = True
    
    # 添加表格
    table = doc_a.add_table(rows=2, cols=3)
    table.style = 'LightShading-Accent1'
    for i in range(2):
        for j in range(3):
            cell = table.cell(i, j)
            cell.text = f'行{i+1}列{j+1}'
    
    doc_a.add_paragraph('这是最后一段内容。')
    doc_a.save('source_document.docx')
    print("源文档已创建: source_document.docx")
    
    # 创建目标文档B
    doc_b = Document()
    doc_b.add_heading('目标文档B', 0)
    doc_b.add_paragraph('这是目标文档的开头部分。')
    doc_b.add_paragraph('')
    doc_b.add_paragraph('下面将在这里插入源文档的内容：{{ place1 }}')
    doc_b.add_paragraph('')
    doc_b.add_paragraph('这是插入后的内容。')
    
    # 添加表格测试
    table_b = doc_b.add_table(rows=2, cols=2)
    table_b.cell(0, 0).text = '表格单元格1'
    table_b.cell(0, 1).text = '表格中占位符: {{ place1 }}'
    table_b.cell(1, 0).text = '单元格3'
    table_b.cell(1, 1).text = '单元格4'
    
    doc_b.add_paragraph('文档结束。')
    doc_b.save('target_document.docx')
    print("目标文档已创建: target_document.docx")


def docx_merge():
    """测试文档合并器"""
    print("=" * 50)
    print("测试文档合并工具")
    print("=" * 50)
    
    # 创建测试文档
    create_test_documents()
    
    # 测试文档合并
    print("\n开始合并文档...")
    merger = DocxContentMerger()
    
    try:
        result_file = merger.merge_documents(
            source_file='source_document.docx',
            target_file='target_document.docx',
            output_file='merged_result.docx'
        )
        
        # 验证结果
        result_doc = Document(result_file)
        print(f"\n合并结果:")
        print(f"输出文件: {result_file}")
        print(f"文档段落数: {len(result_doc.paragraphs)}")
        print(f"文档表格数: {len(result_doc.tables)}")
        
        # 检查占位符是否被替换
        has_placeholder = False
        for para in result_doc.paragraphs:
            if '{{ place1 }}' in para.text:
                has_placeholder = True
                break
        
        if not has_placeholder:
            print("✓ 占位符已成功替换")
        else:
            print("✗ 占位符未被完全替换")
        
        return True
        
    except Exception as e:
        print(f"错误: {e}")
        return False
    finally:
        # 清理测试文件
        for file in ['source_document.docx', 'target_document.docx', 'merged_result.docx']:
            if os.path.exists(file):
                os.remove(file)
                print(f"已清理: {file}")


def test_merge_docx():
    """主函数"""
    print("Word文档合并工具")
    print("=" * 50)
    
    # 运行测试
    if docx_merge():
        print("\n测试完成!")
    else:
        print("\n测试失败!")
    
    print("\n使用示例:")
    print("1. 基本使用:")
    print("   merger = DocxContentMerger()")
    print("   result = merger.merge_documents('source.docx', 'target.docx', 'output.docx')")
    print("")
    print("2. 自定义占位符:")
    print("   result = merger.merge_with_custom_placeholder('source.docx', 'target.docx', 'my_placeholder', 'output.docx')")
    print("")
    print("注意事项:")
    print("- 确保目标文档中包含占位符，如: {{ place1 }}")
    print("- 源文档的所有内容将被插入到占位符位置")
    print("- 支持段落和表格的内容复制")