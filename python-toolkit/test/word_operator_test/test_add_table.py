from docx import Document
from docx.table import Table
from docx.enum.text import WD_ALIGN_PARAGRAPH

import sys
sys.path.append('../src')
from word import word_operator

def test_add_table_example():
    doc = Document()
    # 添加表格：4列，初始行数3行
    table = doc.add_table(rows=3, cols=4)
    
    # ---------- 关键：为整个表格添加内外边框 ----------
    # 使用内置的“Table Grid”样式，该样式包含所有单元格边框
    table.style = 'Table Grid'
    
    # 第一行：合并 A 到 D 列
    row0_cells = table.rows[0].cells
    merged_cell = row0_cells[0].merge(row0_cells[3])  # 合并从索引0到3
    merged_cell.text = "这是合并单元格中的一行文字"
    # 可以设置居中
    merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 第二行：表头，加粗，背景灰色
    row1_cells = table.rows[1].cells
    headers = ["列A", "列B", "列C", "列D"]
    for i, cell in enumerate(row1_cells):
        cell.text = headers[i]
        word_operator.set_cell_text_bold(cell)
        word_operator.set_cell_background(cell, (191, 191, 191))  # 浅灰色
    
    # 第三行：示例数据
    row2_cells = table.rows[2].cells
    data = ["数据A1", "数据B1", "数据C1", "数据D1"]
    for i, cell in enumerate(row2_cells):
        cell.text = data[i]
    
    # 保存文档
    doc.save("example_table.docx")


def create_styled_table(doc: Document) -> Table:
    """
    在文档 doc 中创建一个样式化表格（带合并单元格、表头加粗灰底、全边框）
    注意：表格会被添加到文档末尾，但之后可通过 insert_table_at_placeholder 移动
    """
    # 创建一个 3 行 4 列的表格
    table = doc.add_table(rows=3, cols=4)
    
    # 设置表格样式为“网格型”，使所有单元格带边框
    table.style = 'Table Grid'
    
    # 第一行：合并单元格
    row0_cells = table.rows[0].cells
    merged_cell = row0_cells[0].merge(row0_cells[3])
    merged_cell.text = "这是合并单元格中的一行文字"
    merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 第二行：表头（加粗+灰色背景）
    headers = ["列A", "列B", "列C", "列D"]
    row1_cells = table.rows[1].cells
    for i, cell in enumerate(row1_cells):
        cell.text = headers[i]
        word_operator.set_cell_text_bold(cell)
        word_operator.set_cell_background(cell, (191, 191, 191))   # 浅灰色
    
    # 第三行：示例数据
    row2_cells = table.rows[2].cells
    sample_data = ["数据A1", "数据B1", "数据C1", "数据D1"]
    for i, cell in enumerate(row2_cells):
        cell.text = sample_data[i]
    
    return table

def test_insert_table_placeholder_example():
    # 1. 创建一个临时文档，并写入占位符
    doc = Document()
    doc.add_paragraph("这是文档开头的文字。")
    doc.add_paragraph("下面将插入表格：")
    doc.add_paragraph("{{ 1 }}")          # 占位符单独一行
    doc.add_paragraph("这是表格后面的文字。")
    
    # 保存临时文档（可选，仅用于查看占位符）
    temp_path = "temp_with_placeholder.docx"
    doc.save(temp_path)
    print(f"已创建带占位符的临时文档: {temp_path}")
    
    # 2. 重新加载文档（或继续使用当前 doc，但为了演示移动，我们重新加载）
    doc = Document(temp_path)
    
    # 3. 创建样式化表格（注意：此时表格会被添加到文档末尾）
    table = create_styled_table(doc)
    
    # 4. 将表格插入到占位符 "{{ 1 }}" 的位置（注意占位符文本需完全匹配）
    word_operator.insert_table_at_placeholder(doc, "{{ 1 }}", table)
    
    # 5. 保存最终文档
    output_path = "final_with_table.docx"
    doc.save(output_path)
    print(f"已生成最终文档: {output_path}")