import os
from typing import Optional, Dict, Any
from openpyxl import load_workbook
from openpyxl.utils import get_column_letter
from docx import Document
from docx.table import Table, _Cell
from docx.oxml.ns import qn
from docx.oxml import parse_xml
import re


class ExcelToWordTable:
    """
    将Excel表格插入Word文档指定位置的工具类
    """
    
    def __init__(self, excel_path: str, word_path: str = None):
        """
        初始化工具
        
        Args:
            excel_path: Excel文件路径
            word_path: Word文档路径（可选，可后续创建）
        """
        self.excel_path = excel_path
        self.word_path = word_path
        self.word_doc = None
        
        # 验证Excel文件存在
        if not os.path.exists(excel_path):
            raise FileNotFoundError(f"Excel文件不存在: {excel_path}")
    
    def load_excel(self, sheet_name: Optional[str] = None, 
                   start_row: int = 1, start_col: int = 1,
                   end_row: Optional[int] = None, 
                   end_col: Optional[int] = None) -> list:
        """
        从Excel加载表格数据
        
        Args:
            sheet_name: 工作表名，默认第一个工作表
            start_row: 起始行（1-based）
            start_col: 起始列（1-based）
            end_row: 结束行，None表示到最后一行
            end_col: 结束列，None表示到最后一列
        
        Returns:
            list: 二维列表的表格数据
        """
        # 加载Excel工作簿
        wb = load_workbook(self.excel_path, data_only=True)  # data_only=True获取计算后的值
        
        # 获取工作表
        if sheet_name:
            ws = wb[sheet_name]
        else:
            ws = wb.active
        
        # 确定读取范围
        max_row = end_row if end_row else ws.max_row
        max_col = end_col if end_col else ws.max_column
        
        # 确保范围有效
        if start_row > max_row or start_col > max_col:
            return []
        
        # 读取数据
        table_data = []
        for row in ws.iter_rows(min_row=start_row, max_row=max_row,
                               min_col=start_col, max_col=max_col,
                               values_only=True):
            # 将None转换为空字符串
            row_data = ["" if cell is None else str(cell) for cell in row]
            table_data.append(row_data)
        
        wb.close()
        return table_data
    
    def create_word_document(self, template_path: Optional[str] = None):
        """
        创建或加载Word文档
        
        Args:
            template_path: 模板文档路径（可选）
        """
        if template_path:
            if not os.path.exists(template_path):
                raise FileNotFoundError(f"Word模板文件不存在: {template_path}")
            self.word_doc = Document(template_path)
        else:
            self.word_doc = Document()
    
    def save_word_document(self, output_path: Optional[str] = None):
        """
        保存Word文档
        
        Args:
            output_path: 输出路径，不指定则使用初始化时的路径
        """
        if not self.word_doc:
            raise ValueError("Word文档未创建，请先调用create_word_document()")
        
        save_path = output_path or self.word_path
        if not save_path:
            raise ValueError("未指定保存路径")
        
        self.word_doc.save(save_path)
        print(f"Word文档已保存: {save_path}")
    
    def insert_table_at_placeholder(self, table_data: list, 
                                  placeholder: str = "{{ place1 }}",
                                  auto_adjust_columns: bool = True,
                                  table_style: Optional[str] = "Table Grid"):
        """
        在指定占位符位置插入表格
        
        Args:
            table_data: 表格数据（二维列表）
            placeholder: 占位符文本，如 {{ place1 }}
            auto_adjust_columns: 是否自动调整列宽
            table_style: 表格样式
        """
        if not self.word_doc:
            raise ValueError("Word文档未创建，请先调用create_word_document()")
        
        if not table_data:
            print("警告: 表格数据为空")
            return
        
        # 在文档中搜索占位符
        found = False
        for paragraph in self.word_doc.paragraphs:
            if placeholder in paragraph.text:
                found = True
                # 创建表格
                rows = len(table_data)
                cols = len(table_data[0]) if rows > 0 else 0
                
                # 在段落前插入表格
                table = self.word_doc.add_table(rows=rows, cols=cols, style=table_style)
                
                # 填充表格数据
                for i, row_data in enumerate(table_data):
                    for j, cell_data in enumerate(row_data):
                        cell = table.cell(i, j)
                        cell.text = str(cell_data)
                
                # 自动调整列宽
                if auto_adjust_columns:
                    self._auto_adjust_column_width(table)
                
                # 替换占位符
                paragraph.text = paragraph.text.replace(placeholder, "")
                
                # 将表格移动到占位符位置
                paragraph._p.addprevious(table._tbl)
                
                # 删除原始空段落
                if paragraph.text.strip() == "":
                    paragraph._p.getparent().remove(paragraph._p)
                
                break
        
        if not found:
            # 如果没有找到占位符，在文档末尾添加表格
            print(f"未找到占位符 '{placeholder}'，表格将添加到文档末尾")
            self.add_table_to_end(table_data, auto_adjust_columns, table_style)
    
    def _auto_adjust_column_width(self, table: Table):
        """自动调整列宽"""
        # 获取每列的最大宽度
        col_widths = [0] * len(table.columns)
        
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                # 计算文本长度（简单估算）
                text_length = len(cell.text)
                if text_length > col_widths[j]:
                    col_widths[j] = text_length
        
        # 设置列宽（这里使用相对单位）
        for j, width in enumerate(col_widths):
            # 转换为厘米（简单估算）
            width_cm = min(width * 0.3, 8)  # 最大8cm
            for cell in table.columns[j].cells:
                cell.width = width_cm
    
    def add_table_to_end(self, table_data: list, 
                        auto_adjust_columns: bool = True,
                        table_style: Optional[str] = "Table Grid"):
        """
        在文档末尾添加表格
        
        Args:
            table_data: 表格数据
            auto_adjust_columns: 是否自动调整列宽
            table_style: 表格样式
        """
        if not table_data:
            return
        
        rows = len(table_data)
        cols = len(table_data[0]) if rows > 0 else 0
        
        # 添加表格
        table = self.word_doc.add_table(rows=rows, cols=cols, style=table_style)
        
        # 填充数据
        for i, row_data in enumerate(table_data):
            for j, cell_data in enumerate(row_data):
                cell = table.cell(i, j)
                cell.text = str(cell_data)
        
        # 自动调整列宽
        if auto_adjust_columns:
            self._auto_adjust_column_width(table)
    
    def replace_multiple_placeholders(self, placeholder_data_dict: Dict[str, list],
                                     auto_adjust_columns: bool = True,
                                     table_style: Optional[str] = "Table Grid"):
        """
        批量替换多个占位符
        
        Args:
            placeholder_data_dict: 占位符和表格数据的字典
            auto_adjust_columns: 是否自动调整列宽
            table_style: 表格样式
        """
        for placeholder, table_data in placeholder_data_dict.items():
            self.insert_table_at_placeholder(
                table_data, placeholder, auto_adjust_columns, table_style
            )
