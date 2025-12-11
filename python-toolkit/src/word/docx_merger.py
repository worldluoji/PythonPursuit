from docx import Document
import re
import copy
import os
from typing import List, Tuple, Optional


class DocxContentMerger:
    """
    Word文档合并工具
    功能：读取文件A的所有内容，插入到文件B中指定占位符位置
    占位符格式：{{ placeholder_name }}
    """
    
    def __init__(self, placeholder_format: str = r'\{\{\s*(\w+)\s*\}\}'):
        """
        初始化文档合并器
        
        参数:
        - placeholder_format: 占位符的正则表达式模式
        """
        self.placeholder_pattern = placeholder_format
    
    def extract_all_content(self, file_path: str) -> List[Tuple[str, any]]:
        """
        从源文档提取所有内容（段落、表格）
        
        参数:
        - file_path: 源文档路径
        
        返回:
        - List[Tuple[str, any]]: 内容元素列表，每个元素是(类型, 内容)元组
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        try:
            source_doc = Document(file_path)
            content_elements = []
            
            # 按顺序读取所有段落
            for para in source_doc.paragraphs:
                if para.text.strip():  # 只添加非空段落
                    content_elements.append(('paragraph', self._copy_paragraph(para)))
            
            # 按顺序读取所有表格
            for table in source_doc.tables:
                content_elements.append(('table', self._copy_table(table)))
            
            return content_elements
            
        except Exception as e:
            raise Exception(f"读取文档失败: {e}")
    
    def _copy_paragraph(self, source_para) -> Document:
        """
        深度复制段落及其格式
        
        参数:
        - source_para: 源段落对象
        
        返回:
        - Document: 包含复制段落的新文档对象
        """
        # 创建新文档
        temp_doc = Document()
        
        # 创建新段落
        new_para = temp_doc.add_paragraph()
        
        # 复制段落样式
        if source_para.style:
            new_para.style = source_para.style.name
        
        # 复制段落格式 - 使用安全的属性访问
        try:
            new_para.paragraph_format.alignment = source_para.paragraph_format.alignment
            new_para.paragraph_format.left_indent = source_para.paragraph_format.left_indent
            new_para.paragraph_format.right_indent = source_para.paragraph_format.right_indent
            new_para.paragraph_format.first_line_indent = source_para.paragraph_format.first_line_indent
            new_para.paragraph_format.line_spacing = source_para.paragraph_format.line_spacing  # 修复：line_sp -> line_spacing
            new_para.paragraph_format.space_before = source_para.paragraph_format.space_before
            new_para.paragraph_format.space_after = source_para.paragraph_format.space_after
        except AttributeError as e:
            # 如果某些属性不存在，跳过它们
            pass
        
        # 复制所有runs及其格式
        for run in source_para.runs:
            new_run = new_para.add_run(run.text)
            
            # 复制字体格式
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline
            
            # 安全地复制字体名称
            if hasattr(run.font, 'name') and run.font.name:
                new_run.font.name = run.font.name
            
            # 安全地复制字体大小
            if hasattr(run.font, 'size') and run.font.size:
                new_run.font.size = run.font.size
            
            # 安全地复制字体颜色
            if hasattr(run.font, 'color') and run.font.color and hasattr(run.font.color, 'rgb') and run.font.color.rgb:
                new_run.font.color.rgb = run.font.color.rgb
        
        return temp_doc
    
    def _copy_table(self, source_table) -> Document:
        """
        深度复制表格
        
        参数:
        - source_table: 源表格对象
        
        返回:
        - Document: 包含复制表格的新文档对象
        """
        # 创建新文档
        temp_doc = Document()
        
        # 获取表格的行列数
        rows = len(source_table.rows)
        cols = 0
        if rows > 0:
            # 获取第一行的列数
            cols = len(source_table.rows[0].cells)
        
        if rows > 0 and cols > 0:
            # 创建新表格
            new_table = temp_doc.add_table(rows=rows, cols=cols)
            
            # 复制表格内容和基本格式
            for i in range(rows):
                for j in range(cols):
                    if i < len(source_table.rows) and j < len(source_table.rows[i].cells):
                        source_cell = source_table.cell(i, j)
                        new_cell = new_table.cell(i, j)
                        new_cell.text = source_cell.text
        
        return temp_doc
    
    def find_placeholder_positions(self, doc: Document) -> dict:
        """
        在文档中查找所有占位符的位置
        
        参数:
        - doc: Document对象
        
        返回:
        - dict: 占位符位置信息
        """
        positions = {
            'paragraphs': [],  # 在段落中找到的占位符
            'tables': []       # 在表格中找到的占位符
        }
        
        # 在段落中查找
        for para_idx, paragraph in enumerate(doc.paragraphs):
            matches = re.findall(self.placeholder_pattern, paragraph.text)
            if matches:
                for match in matches:
                    positions['paragraphs'].append({
                        'placeholder': match,
                        'paragraph_idx': para_idx,
                        'paragraph': paragraph,
                        'full_text': paragraph.text
                    })
        
        # 在表格中查找
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for para_in_cell_idx, cell_para in enumerate(cell.paragraphs):
                        matches = re.findall(self.placeholder_pattern, cell_para.text)
                        if matches:
                            for match in matches:
                                positions['tables'].append({
                                    'placeholder': match,
                                    'table_idx': table_idx,
                                    'row_idx': row_idx,
                                    'cell_idx': cell_idx,
                                    'para_in_cell_idx': para_in_cell_idx,
                                    'paragraph': cell_para,
                                    'full_text': cell_para.text
                                })
        
        return positions
    
    def replace_in_paragraph(self, doc: Document, placeholder_info: dict, source_content: List[Tuple[str, any]]) -> bool:
        """
        在段落中替换占位符
        
        参数:
        - doc: 目标文档对象
        - placeholder_info: 占位符位置信息
        - source_content: 源内容列表
        
        返回:
        - bool: 是否成功替换
        """
        try:
            paragraph = placeholder_info['paragraph']
            placeholder = placeholder_info['placeholder']
            placeholder_text = f"{{{{ {placeholder} }}}}"
            
            # 如果段落中不包含占位符，直接返回
            if placeholder_text not in paragraph.text:
                return False
            
            print(f"在段落中替换占位符: {placeholder_text}")
            
            # 获取段落元素
            p = paragraph._element
            p_parent = p.getparent()
            p_index = list(p_parent).index(p)
            
            # 在占位符位置插入源内容
            for content_type, content_doc in source_content:
                if content_type == 'paragraph':
                    # 复制段落
                    source_para = content_doc.paragraphs[0]
                    
                    # 创建新段落元素
                    new_p = copy.deepcopy(source_para._element)
                    
                    # 插入到占位符段落之前
                    p_parent.insert(p_index, new_p)
                    p_index += 1
                    
                elif content_type == 'table':
                    # 复制表格
                    if content_doc.tables:
                        source_table = content_doc.tables[0]
                        
                        # 创建新表格元素
                        new_tbl = copy.deepcopy(source_table._element)
                        
                        # 插入到占位符段落之前
                        p_parent.insert(p_index, new_tbl)
                        p_index += 1
            
            # 删除原占位符段落
            p_parent.remove(p)
            
            return True
            
        except Exception as e:
            print(f"段落替换失败: {e}")
            return False
    
    def replace_in_table_cell(self, doc: Document, placeholder_info: dict, source_content: List[Tuple[str, any]]) -> bool:
        """
        在表格单元格中替换占位符
        
        参数:
        - doc: 目标文档对象
        - placeholder_info: 占位符位置信息
        - source_content: 源内容列表
        
        返回:
        - bool: 是否成功替换
        """
        try:
            cell_para = placeholder_info['paragraph']
            placeholder = placeholder_info['placeholder']
            placeholder_text = f"{{{{ {placeholder} }}}}"
            
            # 获取单元格
            cell = cell_para._element.getparent().getparent()  # tc -> tr -> tbl
            
            # 清空单元格内容
            for elem in cell.iterchildren():
                if elem.tag.endswith('p'):  # 段落元素
                    cell.remove(elem)
            
            # 在单元格中插入源内容
            for content_type, content_doc in source_content:
                if content_type == 'paragraph':
                    # 获取源段落
                    source_para = content_doc.paragraphs[0]
                    
                    # 复制段落元素到单元格
                    new_p = copy.deepcopy(source_para._element)
                    cell.append(new_p)
                    
                elif content_type == 'table':
                    # 在单元格中不能直接插入表格，用段落代替
                    if content_doc.tables:
                        source_table = content_doc.tables[0]
                        for row in source_table.rows:
                            row_text = ' | '.join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                            if row_text:
                                # 在单元格中添加一个段落
                                new_para = Document().add_paragraph(row_text)
                                new_p = copy.deepcopy(new_para._element)
                                cell.append(new_p)
            
            return True
            
        except Exception as e:
            print(f"表格单元格替换失败: {e}")
            return False
    
    def merge_documents(self, source_file: str, target_file: str, 
                       output_file: Optional[str] = None) -> str:
        """
        合并两个文档：将源文档内容插入到目标文档的占位符位置
        
        参数:
        - source_file: 源文档路径
        - target_file: 目标文档路径
        - output_file: 输出文件路径，默认为'merged_document.docx'
        
        返回:
        - str: 输出文件路径
        """
        if output_file is None:
            output_file = 'merged_document.docx'
        
        # 检查文件是否存在
        for file_path in [source_file, target_file]:
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"文件不存在: {file_path}")
        
        print(f"正在合并文档...")
        print(f"源文档: {source_file}")
        print(f"目标文档: {target_file}")
        
        try:
            # 1. 提取源文档内容
            print("正在读取源文档内容...")
            source_content = self.extract_all_content(source_file)
            print(f"提取到 {len(source_content)} 个内容元素")
            
            # 2. 读取目标文档
            target_doc = Document(target_file)
            
            # 3. 查找所有占位符位置
            placeholder_positions = self.find_placeholder_positions(target_doc)
            all_placeholders = placeholder_positions['paragraphs'] + placeholder_positions['tables']
            
            if not all_placeholders:
                print("警告: 未找到任何占位符，将在文档末尾添加源文档内容")
                # 在文档末尾添加源内容
                for content_type, content_doc in source_content:
                    if content_type == 'paragraph':
                        # 添加段落
                        source_para = content_doc.paragraphs[0]
                        new_para = target_doc.add_paragraph()
                        self._copy_paragraph_to_doc(source_para, target_doc, new_para)
                    elif content_type == 'table':
                        # 添加表格
                        if content_doc.tables:
                            self._copy_table_to_doc(content_doc.tables[0], target_doc)
            
            else:
                # 4. 按位置排序（从后往前替换，避免索引变化）
                all_placeholders.sort(key=lambda x: (
                    x.get('paragraph_idx', -1), 
                    x.get('table_idx', -1), 
                    x.get('row_idx', -1), 
                    x.get('cell_idx', -1)
                ), reverse=True)
                
                # 5. 替换所有占位符
                replaced_count = 0
                for placeholder_info in all_placeholders:
                    if 'paragraph_idx' in placeholder_info:
                        # 在段落中替换
                        if self.replace_in_paragraph(target_doc, placeholder_info, source_content):
                            replaced_count += 1
                    else:
                        # 在表格单元格中替换
                        if self.replace_in_table_cell(target_doc, placeholder_info, source_content):
                            replaced_count += 1
                
                print(f"成功替换 {replaced_count} 个占位符")
            
            # 6. 保存结果
            target_doc.save(output_file)
            print(f"文档合并完成，已保存为: {output_file}")
            
            return output_file
            
        except Exception as e:
            raise Exception(f"文档合并失败: {e}")
    
    def _copy_paragraph_to_doc(self, source_para, target_doc, target_para):
        """复制段落内容到目标文档的段落"""
        # 清空目标段落
        target_para.clear()
        
        # 复制所有runs
        for run in source_para.runs:
            new_run = target_para.add_run(run.text)
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline
            if hasattr(run.font, 'name') and run.font.name:
                new_run.font.name = run.font.name
            if hasattr(run.font, 'size') and run.font.size:
                new_run.font.size = run.font.size
    
    def _copy_table_to_doc(self, source_table, target_doc):
        """复制表格到目标文档"""
        rows = len(source_table.rows)
        cols = 0
        if rows > 0:
            cols = len(source_table.rows[0].cells)
        
        if rows > 0 and cols > 0:
            new_table = target_doc.add_table(rows=rows, cols=cols)
            for i in range(rows):
                for j in range(cols):
                    if i < len(source_table.rows) and j < len(source_table.rows[i].cells):
                        source_cell = source_table.cell(i, j)
                        new_cell = new_table.cell(i, j)
                        new_cell.text = source_cell.text
    
    def merge_with_custom_placeholder(self, source_file: str, target_file: str, 
                                     placeholder: str, output_file: Optional[str] = None) -> str:
        """
        使用自定义占位符合并文档
        
        参数:
        - source_file: 源文档路径
        - target_file: 目标文档路径
        - placeholder: 占位符名称，如"place1"
        - output_file: 输出文件路径
        
        返回:
        - str: 输出文件路径
        """
        # 临时修改占位符模式
        original_pattern = self.placeholder_pattern
        self.placeholder_pattern = rf'\{{{{\s*{placeholder}\s*}}}}'
        
        try:
            result = self.merge_documents(source_file, target_file, output_file)
            return result
        finally:
            # 恢复原始模式
            self.placeholder_pattern = original_pattern


